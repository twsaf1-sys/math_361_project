from docx import Document
from docx.shared import Inches

# Create a new Word document
document = Document()

# Add a heading
document.add_heading('My Python Code Example', level=1)

# Add the Python code as a paragraph (you might need to apply a specific style for code)
code_snippet = """
def factorial(n):
    if n == 0:
        return 1
    else:
        return n * factorial(n-1)

result = factorial(5)
print(f"The factorial of 5 is: {result}")
"""
document.add_paragraph(code_snippet, style='Code') # Assuming you have a 'Code' style defined in Word

# Save the document
document.save('python_code_document.docx')
